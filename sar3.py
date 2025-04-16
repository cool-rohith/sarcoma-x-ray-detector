import os
from datetime import datetime
import streamlit as st
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
import google.generativeai as genai

# Set API Key for Google Gemini AI
api_key = "AIzaSyAP0pnc4IQhWLxU4-9G6RIshrgG21IiQck"
if not api_key:
    raise ValueError("Google Gemini API key is not set.")

# Configure Gemini
genai.configure(api_key=api_key)

# AI Model Configuration
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)

# Analyze X-ray image using Gemini AI
def analyze_xray_image_with_api(image_file):
    image = Image.open(image_file).convert("RGB")

    current_date = datetime.now().strftime('%Y-%m-%d')

    prompt = f"""
Analyze this X-ray image and generate a detailed medical report focusing on *sarcoma detection*. The report should include:
1. *Possible signs of sarcoma* (bone/soft tissue abnormalities, tumor growth, calcifications).
2. *Risk assessment* (low, medium, high).
3. *Possible differential diagnoses* (other conditions that may mimic sarcoma).
4. *Recommended next steps* (further imaging, biopsy, specialist consultation).
5. *Caution and disclaimer* for AI-based analysis.
6. *Reporting date:* {current_date}.
"""

    chat = model.start_chat()
    response = chat.send_message([prompt, image])  # pass PIL.Image directly
    return response.text

# Create Word report
def create_xray_report(report_text, xray_image):
    doc = Document()
    doc.add_heading('Sarcoma Detection Report (X-ray)', 0)

    for line in report_text.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading('X-ray Image:', level=1)
    image_stream = BytesIO(xray_image.getvalue())
    doc.add_picture(image_stream, width=Inches(6))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Streamlit App
def main():
    st.title("🩻 AI-Based Sarcoma Detection from X-ray")
    st.markdown("Upload an X-ray image and let AI generate a detailed sarcoma detection report.")

    xray_image = st.file_uploader("Upload X-ray Image (JPG or PNG)", type=["jpg", "png"])

    if xray_image:
        st.image(xray_image, caption="Uploaded X-ray", use_column_width=True)

        if st.button("🔍 Generate Report"):
            with st.spinner("Analyzing the X-ray image..."):
                try:
                    report = analyze_xray_image_with_api(xray_image)
                    st.success("Analysis Complete!")
                    st.header("📝 Report")
                    st.markdown(report)

                    doc_file = create_xray_report(report, xray_image)
                    st.download_button(
                        label="📄 Download Report",
                        data=doc_file,
                        file_name="Sarcoma_Xray_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Something went wrong: {e}")

if __name__ == '__main__':
    main()
